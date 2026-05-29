from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# 页面设置
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

def add_title(text):
    p = doc.add_heading(text, level=0)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def add_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def add_h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run.bold = bold
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(9.5)
                run.bold = True
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9.5)
    doc.add_paragraph()  # 表后空行

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)

# ====== 开始写内容 ======

add_title('抖店上架助手 — 产品介绍')

# 解决的问题
add_h2('解决的问题')
add_para('在抖店上架服装商品，每件商品要在平台上填几十个字段，大量重复劳动。这个工具把上架流程变成：放图片 → 核对表格 → 点确认，其余自动完成。')

# 工作流程
add_h2('工作流程')
add_code('放入商品图片 → AI识图提取属性 → 自动生成35字标题 → 写入分区Excel表格')
add_code('                                                    ↓')
add_code('                                   人工核对/修改 → 确认后一键自动填表上架')

# 核心功能
add_h2('核心功能')
add_table(
    ['模块', '做什么', '用到的技术'],
    [
        ['AI 识图', '从服装图片识别领型、袖长、颜色、风格等属性', '通义千问 VL（免费，2000次/天）'],
        ['标题生成', '按抖音风格自动生成约35字商品标题', 'AI 大模型'],
        ['分区表格', '6个类目各自独立的Sheet，只显示相关字段', 'Excel（openpyxl）'],
        ['浏览器自动化', '打开抖店后台，逐步填写所有字段', 'Playwright'],
    ]
)

# 支持的类目
add_h2('支持的类目')
add_para('连衣裙 / 上衣 / 裤子 / 半裙 / 外套 / 套装（6个Sheet，各填各的）')

# 输入表格长什么样
add_h2('输入表格长什么样')
add_para('一个 Excel 文件，包含 7 个 Sheet，打开后一目了然：')

add_table(
    ['Sheet', '内容', '说明'],
    [
        ['Sheet 1', '通用配置', '品牌/发货地/运费/退货… 设一次就不用再管'],
        ['Sheet 2', '连衣裙', '标题 | 领型 | 袖长 | 裙长 | 裙型 | 版型 | 风格 | 颜色…'],
        ['Sheet 3', '上衣', '标题 | 领型 | 袖长 | 衣长 | 版型 | 衣门襟 | 风格…'],
        ['Sheet 4', '裤子', '标题 | 裤长 | 裤型 | 腰型 | 厚薄 | 风格…'],
        ['Sheet 5', '半裙', '标题 | 裙长 | 裙型 | 腰型 | 版型 | 风格…'],
        ['Sheet 6', '外套', '标题 | 领型 | 袖长 | 衣长 | 厚薄 | 衣门襟…'],
        ['Sheet 7', '套装', '标题 | 上装领型 | 上装袖长 | 下装类型 | 下装长度…'],
    ]
)

add_h3('每一行 = 一个商品')
add_bullet('第 1 行：示例数据（供参考格式）')
add_bullet('第 2 行起：实际商品，每行一个')
add_bullet('同类目的商品放在同一个 Sheet，互不干扰')
add_bullet('不同类目只显示自己需要的列，不相关的字段不出现')

add_h3('三类列的颜色区分')
add_table(
    ['颜色', '含义', '举例'],
    [
        ['蓝色列', 'AI 自动填，人工可改', '领型、袖长、颜色、风格…'],
        ['黄色列', '全人工填写', '售价、库存、图片路径'],
        ['绿色列', '固定值，设完不改', '品牌、发货地、面料材质…'],
    ]
)

# 自动化覆盖范围
add_h2('自动化覆盖范围')
add_table(
    ['页面区域', '填写内容', '来源'],
    [
        ['上传页', '5张1:1主图', '人工裁剪好的图片文件'],
        ['上传页', '商品标题', 'AI生成'],
        ['基础信息', '导购短标题(12字)', 'AI凝练'],
        ['基础信息', '面料材质', '固定：聚酯纤维85%+其他15%'],
        ['基础信息', '品牌', '固定：无品牌'],
        ['基础信息', '类目特定属性', 'AI识图 → 匹配下拉选项'],
        ['图文信息', '主图(1:1)×5、主图(3:4)×5、详情', '自动填入'],
        ['价格库存', '发货模式、发货时间', '按配置填入'],
        ['价格库存', 'SKU创建+规格图', '按表格数据自动填'],
        ['价格库存', '尺码固定值', 'S:80-95 M:95-105 L:105-115 XL:115-125'],
        ['价格库存', '价格、库存', '按表格数据填入'],
        ['价格库存', '订单库存计数', '固定：付款减库存'],
        ['服务与履约', '运费模板、售后、商品状态', '固定值，无需改动'],
    ]
)

# 不自动做的事
add_h2('不自动做的事')
add_bullet('图片裁剪（人工更可靠）')
add_bullet('最终点"发布"（人工审核后手动发布，发布前商品状态为"下架"）')

# 理想效果
add_h2('理想效果')
add_para('每件商品从上架到发布仅需 2~3 分钟，且每一步都有表格记录留底。')

# 保存
output = r'C:\Users\86153\Desktop\抖店上架助手\产品介绍.docx'
doc.save(output)
print(f'OK: {output}')
